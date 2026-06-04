"""
Insert §27.1 subsection into the CI/CD rules docx, just before §28.
"""
import copy
from lxml import etree
import docx
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = "/home/ai/.claude/uploads/21225044-61b1-4f0e-930b-fbe3ec689e73/2a254c1d-devops_procedure_deploying_from_github.docx"

doc = Document(DOCX_PATH)

# Locate §28 paragraph index — insert new content just before it
para_28_index = None
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("§28 "):
        para_28_index = i
        break

if para_28_index is None:
    raise RuntimeError("Could not find §28 paragraph")

print(f"§28 is at paragraph index {para_28_index}. Will insert before it.")

# The paragraphs list is a view; we need the XML element of §28 to insert before it.
para_28_elem = doc.paragraphs[para_28_index]._element
parent = para_28_elem.getparent()

new_paragraphs_data = [
    {
        "text": "§27.1 u2giants/twenty — approved fast-path exceptions",
        "bold": True,
    },
    {
        "text": (
            "The u2giants/twenty application (GitHub: u2giants/twenty, deployed to crm.designflow.app) "
            "uses two commit-scope fast paths implemented in .github/workflows/build-and-push.yml. "
            "These are approved exceptions under §27 and §25."
        ),
        "bold": False,
    },
    {
        "text": (
            "Fast path 1 — frontend_fast_path. "
            "Trigger: ALL changed files are under packages/twenty-front/** or packages/twenty-front-component-renderer/**. "
            "What is skipped: the build-and-push-full job (backend compile, backend smoke tests). "
            "What still runs: lint, twenty-front Jest suite, frontend-only Docker build, Coolify deploy. "
            "Rationale: frontend source is isolated from backend source; a UI-only commit cannot alter server behavior."
        ),
        "bold": False,
    },
    {
        "text": (
            "Fast path 2 — backend_only_path. "
            "Trigger: ALL changed files are under packages/twenty-server/** or packages/twenty-emails/**. "
            "What is skipped: the 'Test twenty-front in band' Jest step; the frontend Vite prebuild "
            "(replaced by a GHA actions/cache restore keyed on twenty-front/twenty-ui/twenty-shared/"
            "twenty-front-component-renderer source hashes + yarn.lock). "
            "What still runs: backend lint, backend tests, server smoke test, Docker build, Coolify deploy. "
            "Rationale: a backend-only change cannot alter the frontend bundle; the cache is content-addressed "
            "so a hit proves the bundle is identical to the last verified build."
        ),
        "bold": False,
    },
    {
        "text": (
            "Neither fast path skips the Docker image build or the Coolify deployment trigger. "
            "Any commit touching files outside the fast-path boundaries runs the full pipeline. "
            "Mixed commits (both frontend and backend files changed) always run the full path."
        ),
        "bold": False,
    },
]

def make_paragraph_element(text, bold=False):
    """Create a <w:p> XML element with optional bold run."""
    p = docx.oxml.OxmlElement("w:p")
    r = docx.oxml.OxmlElement("w:r")
    if bold:
        rPr = docx.oxml.OxmlElement("w:rPr")
        b = docx.oxml.OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = docx.oxml.OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p

# Insert all new paragraphs before §28, in order
for item in new_paragraphs_data:
    new_p_elem = make_paragraph_element(item["text"], bold=item["bold"])
    parent.insert(list(parent).index(para_28_elem), new_p_elem)

doc.save(DOCX_PATH)
print("Saved successfully.")

# Verify
doc2 = Document(DOCX_PATH)
for i, p in enumerate(doc2.paragraphs):
    if "§27.1" in p.text or "§28" in p.text:
        print(f"  Para {i}: {p.text[:80]!r}")
